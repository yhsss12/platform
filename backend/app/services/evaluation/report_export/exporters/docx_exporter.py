from __future__ import annotations

from pathlib import Path
from typing import Any

from app.services.evaluation.report_export.export_options import ExportOptions
from app.services.evaluation.report_export.exporters._utils import (
    EPISODE_TABLE_HEADERS,
    VIDEO_TABLE_HEADERS,
    dash,
    episode_table_row,
    metric_display_value,
    metric_status_label,
    ordered_metric_rows,
    section_enabled,
    video_table_row,
)


def export_docx(data: dict[str, Any], out_dir: Path, options: ExportOptions) -> Path:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("导出 Word 需要安装 python-docx：pip install python-docx") from exc

    document = Document()
    meta = data.get("reportMeta") or {}
    document.add_heading("评测报告", level=0)
    document.add_paragraph(f"任务 ID：{dash(meta.get('jobId'))}")
    document.add_paragraph(f"生成时间：{dash(meta.get('generatedAt'))}")
    if data.get("legacyNotice"):
        document.add_paragraph(str(data["legacyNotice"]))

    def add_table_section(title: str, headers: list[str], rows: list[list[str]]) -> None:
        document.add_heading(title, level=1)
        table = document.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        header_cells = table.rows[0].cells
        for index, header in enumerate(headers):
            header_cells[index].text = header
        for row in rows:
            cells = table.add_row().cells
            for index, value in enumerate(row):
                cells[index].text = str(value)

    if section_enabled(options, "basicInfo"):
        basic = data.get("basicInfo") or {}
        add_table_section(
            "1. 基础信息",
            ["字段", "值"],
            [
                ["任务名称", dash(basic.get("taskName"))],
                ["任务 ID", dash(basic.get("jobId"))],
                ["评测类型", dash(basic.get("evaluationTypeLabel"))],
                ["关联任务", dash(basic.get("associatedTaskName"))],
                ["仿真平台", dash(basic.get("simulationPlatform"))],
                ["状态", dash(basic.get("statusLabel"))],
                ["评测对象", dash(basic.get("evaluationObjectLabel"))],
                ["模型资产", dash(basic.get("modelAssetName"))],
                ["机器人", dash(basic.get("robotType"))],
                ["创建时间", dash(basic.get("createdAt"))],
                ["完成时间", dash(basic.get("finishedAt"))],
            ],
        )

    if section_enabled(options, "evaluationConfig"):
        cfg = data.get("evaluationConfig") or {}
        add_table_section(
            "2. 评测配置",
            ["字段", "值"],
            [
                ["Episodes", dash(cfg.get("episodes"))],
                ["Horizon", dash(cfg.get("horizon"))],
                ["Seed", dash(cfg.get("seed"))],
                ["录制视频", dash(cfg.get("recordVideo"))],
                ["任务类型", dash(cfg.get("taskType"))],
                ["模型资产", dash(cfg.get("modelAssetName"))],
            ],
        )

    if section_enabled(options, "metricResults"):
        add_table_section(
            "3. 指标结果",
            ["指标", "值", "单位", "状态", "来源", "说明"],
            [
                [
                    dash(entry.get("displayName") or entry.get("metricId")),
                    metric_display_value(entry),
                    dash(entry.get("unit")),
                    metric_status_label(entry, include_reason=options.include_unavailable_metric_reasons),
                    dash(entry.get("source")),
                    dash(entry.get("reason") if not entry.get("available") else entry.get("description")),
                ]
                for entry in ordered_metric_rows(data)
            ],
        )

    if section_enabled(options, "episodeResults"):
        add_table_section(
            "4. Episode 明细",
            EPISODE_TABLE_HEADERS,
            [episode_table_row(row) for row in data.get("episodeResults") or []],
        )

    if section_enabled(options, "videoInfo"):
        add_table_section(
            "5. 视频回放信息",
            VIDEO_TABLE_HEADERS,
            [video_table_row(row) for row in data.get("videoInfo") or []],
        )

    if section_enabled(options, "diagnostics"):
        diag = data.get("diagnostics") or {}
        add_table_section(
            "6. 失败诊断",
            ["字段", "值"],
            [
                ["摘要", dash(diag.get("summary"))],
                ["失败原因", dash(diag.get("failureReason"))],
                ["异常信息", dash(diag.get("errorMessage"))],
                ["是否超时", "是" if diag.get("timedOut") else "否"],
                ["进程异常退出", "是" if diag.get("abnormalExit") else "否"],
            ],
        )

    if section_enabled(options, "runtimeFiles"):
        add_table_section(
            "7. 原始文件索引",
            ["文件名", "相对路径", "是否存在", "大小", "说明"],
            [
                [
                    dash(row.get("filename")),
                    dash(row.get("relativePath")),
                    "是" if row.get("exists") else "否",
                    dash(row.get("sizeLabel")),
                    dash(row.get("description")),
                ]
                for row in data.get("runtimeFiles") or []
            ],
        )

    path = out_dir / "report.docx"
    document.save(path)
    return path
